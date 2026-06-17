"""
document_parser.py
Extracts structured content (text sections, tables, images)
from uploaded PDF or DOCX files.
"""
import io
import re
import base64
import logging

logger = logging.getLogger(__name__)

# ── Section keyword patterns ─────────────────────────────────────
SECTION_PATTERNS = [
    r'^\s*(introduction|purpose|scope|background|overview)\s*$',
    r'^\s*(procedure|procedures|method|process|steps|workflow)\s*$',
    r'^\s*(results?|findings?|output|observations?)\s*$',
    r'^\s*(safety|hazards?|precautions?|ppe|warnings?)\s*$',
    r'^\s*(references?|bibliography|appendix)\s*$',
    r'^\s*(conclusion|summary|remarks?)\s*$',
    r'^\s*\d+[\.\)]\s*.{3,60}$',   # numbered section: "1. Introduction"
]


def _looks_like_heading(text: str) -> bool:
    text = text.strip()
    if not text:
        return False
    for pat in SECTION_PATTERNS:
        if re.match(pat, text, re.IGNORECASE):
            return True
    # Short lines in ALL CAPS or Title Case that aren't sentences
    if len(text) < 80 and text == text.title() and not text.endswith('.'):
        return True
    if len(text) < 80 and text == text.upper() and len(text.split()) <= 6:
        return True
    return False


def _build_sections(paragraphs: list) -> list:
    """
    Given a flat list of {'type': 'text'|'table'|'image', ...} dicts,
    group them into named sections.
    """
    sections = []
    current = {'name': 'Introduction', 'components': []}

    for para in paragraphs:
        if para['type'] == 'heading':
            if current['components']:
                sections.append(current)
            current = {'name': para['text'], 'components': []}
        else:
            current['components'].append(para)

    if current['components'] or not sections:
        sections.append(current)

    # Ensure we always have at least one section
    if not sections:
        sections = [{'name': 'Content', 'components': []}]

    return sections


# ── PDF Parser ───────────────────────────────────────────────────
def parse_pdf(file_bytes: bytes) -> dict:
    """
    Returns:
      {
        'title': str,
        'sections': [ { 'name': str, 'components': [ {...} ] } ],
        'stats': { 'text_blocks': int, 'tables': int, 'images': int }
      }
    """
    try:
        import PyPDF2
    except ImportError:
        return _empty_result('PyPDF2 not installed')

    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        paragraphs = []
        title = ''

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text() or ''
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if not title and len(line) > 5 and page_num == 0:
                    title = line
                    continue
                if _looks_like_heading(line):
                    paragraphs.append({'type': 'heading', 'text': line})
                else:
                    # Merge short consecutive lines into paragraph blocks
                    if paragraphs and paragraphs[-1]['type'] == 'text_block':
                        paragraphs[-1]['content'] += ' ' + line
                    else:
                        paragraphs.append({'type': 'text_block', 'content': line})

        # Convert text_block → text component format
        components_flat = []
        for p in paragraphs:
            if p['type'] == 'text_block':
                components_flat.append({'type': 'text', 'content': f'<p>{p["content"]}</p>'})
            elif p['type'] == 'heading':
                components_flat.append({'type': 'heading', 'text': p['text']})

        sections = _build_sections(components_flat)

        stats = {
            'text_blocks': sum(1 for s in sections for c in s['components'] if c['type'] == 'text'),
            'tables': 0,
            'images': 0,
            'pages': len(reader.pages),
        }

        return {'title': title, 'sections': sections, 'stats': stats}

    except Exception as e:
        logger.error(f'PDF parse error: {e}')
        return _empty_result(str(e))


# ── DOCX Parser ──────────────────────────────────────────────────
def parse_docx(file_bytes: bytes) -> dict:
    """
    Returns same structure as parse_pdf.
    Extracts: paragraphs, headings, tables, and embedded images.
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        return _empty_result('python-docx not installed')

    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        title = ''
        table_count = 0
        image_count = 0

        # Extract tables
        table_data = []
        for table in doc.tables:
            table_count += 1
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                table_data.append(rows)

        # Extract paragraphs and headings
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name.lower() if para.style else ''
            is_heading = 'heading' in style_name

            if not title and text:
                title = text
                if is_heading:
                    continue

            if is_heading or _looks_like_heading(text):
                paragraphs.append({'type': 'heading', 'text': text})
            else:
                if paragraphs and paragraphs[-1]['type'] == 'text_block':
                    paragraphs[-1]['content'] += '\n' + text
                else:
                    paragraphs.append({'type': 'text_block', 'content': text})

        # Extract images from document relationships
        images_b64 = []
        try:
            for rel in doc.part.rels.values():
                if 'image' in rel.reltype:
                    image_count += 1
                    img_bytes = rel.target_part.blob
                    b64 = base64.b64encode(img_bytes).decode('utf-8')
                    # Try to detect image type
                    content_type = getattr(rel.target_part, 'content_type', 'image/png')
                    images_b64.append(f'data:{content_type};base64,{b64}')
        except Exception:
            pass

        # Build flat component list
        components_flat = []
        for p in paragraphs:
            if p['type'] == 'text_block':
                # Convert newlines to <br> for HTML
                html = ''.join(f'<p>{line}</p>' for line in p['content'].split('\n') if line.strip())
                components_flat.append({'type': 'text', 'content': html})
            elif p['type'] == 'heading':
                components_flat.append({'type': 'heading', 'text': p['text']})

        sections = _build_sections(components_flat)

        # Inject tables into sections intelligently
        # Add any tables to the last section that isn't empty
        if table_data and sections:
            last = sections[-1]
            for rows in table_data:
                last['components'].append({'type': 'table', 'rows': rows})

        # Inject images
        if images_b64 and sections:
            last = sections[-1]
            for src in images_b64:
                last['components'].append({
                    'type': 'image',
                    'src': src,
                    'altText': 'Extracted image',
                    'caption': '',
                })

        stats = {
            'text_blocks': sum(1 for s in sections for c in s['components'] if c['type'] == 'text'),
            'tables': table_count,
            'images': image_count,
        }

        return {'title': title, 'sections': sections, 'stats': stats}

    except Exception as e:
        logger.error(f'DOCX parse error: {e}')
        return _empty_result(str(e))


def _empty_result(error_msg: str = '') -> dict:
    return {
        'title': '',
        'sections': [{'name': 'Content', 'components': [
            {'type': 'text', 'content': f'<p>Could not extract content. {error_msg}</p>'}
        ]}],
        'stats': {'text_blocks': 0, 'tables': 0, 'images': 0},
        'error': error_msg,
    }


def parse_document(file_bytes: bytes, filename: str) -> dict:
    """Entry point. Detects format from filename extension."""
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    if ext == 'pdf':
        return parse_pdf(file_bytes)
    elif ext in ('docx', 'doc'):
        return parse_docx(file_bytes)
    else:
        return _empty_result(f'Unsupported file type: .{ext}')
