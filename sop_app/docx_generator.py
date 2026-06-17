"""
docx_generator.py
Generates a properly formatted Word (.docx) document for an SOP.
"""
import io
import re
import base64
import logging
from datetime import datetime

logger = logging.getLogger(__name__)


def generate_sop_docx(sop) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return b''

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    BLUE  = RGBColor(0x00, 0x3D, 0xA5)
    RED   = RGBColor(0xE3, 0x06, 0x13)
    DGRAY = RGBColor(0x0D, 0x1B, 0x3E)
    LGRAY = RGBColor(0xF2, 0xF4, 0xF8)
    MGRAY = RGBColor(0x5A, 0x6A, 0x8A)

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def add_horizontal_line(para, color_hex='003DA5', thickness=12):
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(thickness))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── Cover header table ────────────────────────────────────────
    hdr_tbl = doc.add_table(rows=1, cols=2)
    hdr_tbl.style = 'Table Grid'
    hdr_tbl.autofit = False
    hdr_tbl.columns[0].width = Cm(12)
    hdr_tbl.columns[1].width = Cm(5)
    lc = hdr_tbl.cell(0, 0)
    rc = hdr_tbl.cell(0, 1)
    set_cell_bg(lc, '003DA5')
    set_cell_bg(rc, '003DA5')
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = lp.add_run('HPCL — Standard Operating Procedure')
    lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    lr.font.bold = True
    lr.font.size = Pt(10)
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run(sop.sop_number or 'DRAFT')
    rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    rr.font.size = Pt(9)
    doc.add_paragraph()

    # ── Title ─────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tr = title_para.add_run(sop.title)
    tr.font.size = Pt(20)
    tr.font.bold = True
    tr.font.color.rgb = DGRAY
    add_horizontal_line(title_para, 'E30613', 18)

    doc.add_paragraph()

    # ── Metadata table ────────────────────────────────────────────
    meta_rows = [
        ('Refinery',     sop.refinery.name     if sop.refinery     else '—',
         'Department',   sop.department.name   if sop.department   else '—'),
        ('Unit',         sop.unit.name         if sop.unit         else '—',
         'Version',      sop.version),
        ('Prepared By',  sop.prepared_by.get_full_name() if sop.prepared_by else '—',
         'Effective Date', str(sop.effective_date) if sop.effective_date else '—'),
        ('Status',       sop.get_status_display(),
         'Date',         str(sop.submitted_date)),
    ]

    mt = doc.add_table(rows=len(meta_rows), cols=4)
    mt.style = 'Table Grid'
    mt.autofit = False
    mt.columns[0].width = Cm(3.5)
    mt.columns[1].width = Cm(7.5)
    mt.columns[2].width = Cm(3.5)
    mt.columns[3].width = Cm(5.5)

    for ri, row in enumerate(meta_rows):
        cells = mt.rows[ri].cells
        for ci, text in enumerate(row):
            p = cells[ci].paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9)
            if ci in (0, 2):  # label cells
                run.font.bold = True
                run.font.color.rgb = MGRAY
                set_cell_bg(cells[ci], 'E8EEFA')
            else:
                run.font.color.rgb = DGRAY
                set_cell_bg(cells[ci], 'FFFFFF' if ri % 2 == 0 else 'F2F4F8')

    doc.add_paragraph()

    # ── Tags ──────────────────────────────────────────────────────
    if sop.tags:
        tp = doc.add_paragraph()
        tr2 = tp.add_run(f'Tags: {" · ".join(sop.tags)}')
        tr2.font.size = Pt(9)
        tr2.font.color.rgb = MGRAY
        doc.add_paragraph()

    # ── Sections ──────────────────────────────────────────────────
    for section in sop.sections.all():
        # Section heading
        sh = doc.add_heading(section.name, level=1)
        sh.runs[0].font.color.rgb = BLUE
        sh.runs[0].font.size = Pt(13)
        add_horizontal_line(sh, 'E8EEFA', 6)

        for comp in section.components.all():

            # ── Text ──────────────────────────────────────────────
            if comp.type == 'text' and comp.content:
                html = comp.content
                # Split on <p>, <br>, newlines
                chunks = re.split(r'<br\s*/?>', html)
                paragraphs_html = []
                for chunk in chunks:
                    paras = re.split(r'</p>\s*<p[^>]*>', chunk)
                    paragraphs_html.extend(paras)

                for ph in paragraphs_html:
                    # Strip remaining HTML
                    clean = re.sub(r'<[^>]+>', '', ph).strip()
                    if not clean:
                        continue
                    # Check if it's a numbered step
                    step_match = re.match(r'^(\d+)[\.:\)]\s+(.+)', clean)
                    if step_match:
                        p = doc.add_paragraph(style='List Number')
                        run = p.add_run(step_match.group(2))
                    else:
                        p = doc.add_paragraph()
                        run = p.add_run(clean)
                    run.font.size = Pt(10)
                    run.font.color.rgb = DGRAY

            # ── Numbered Steps (step type) ────────────────────────
            elif comp.type == 'step':
                steps = comp.content.split('\n') if comp.content else []
                for i, step_text in enumerate(steps):
                    if step_text.strip():
                        p = doc.add_paragraph(style='List Number')
                        run = p.add_run(step_text.strip())
                        run.font.size = Pt(10)
                        run.font.color.rgb = DGRAY

            # ── Table ─────────────────────────────────────────────
            elif comp.type == 'table' and comp.table_rows:
                rows = comp.table_rows
                if not rows:
                    continue
                ncols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=ncols)
                tbl.style = 'Table Grid'
                tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

                for ri, row in enumerate(rows):
                    for ci, cell_text in enumerate(row):
                        if ci >= ncols:
                            break
                        cell = tbl.cell(ri, ci)
                        p = cell.paragraphs[0]
                        run = p.add_run(str(cell_text))
                        run.font.size = Pt(9)
                        if ri == 0:
                            run.font.bold = True
                            run.font.color.rgb = DGRAY
                            set_cell_bg(cell, 'E8EEFA')
                        else:
                            run.font.color.rgb = DGRAY
                            if ri % 2 == 0:
                                set_cell_bg(cell, 'F2F4F8')
                doc.add_paragraph()

            # ── Image / Chart ──────────────────────────────────────
            elif comp.type in ('image', 'chart') and comp.image_url:
                try:
                    src = comp.image_url
                    if src.startswith('data:'):
                        _, b64data = src.split(',', 1)
                        img_bytes = base64.b64decode(b64data)
                        img_buf = io.BytesIO(img_bytes)
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(img_buf, width=Cm(14))
                    else:
                        p = doc.add_paragraph()
                        run = p.add_run()
                        run.add_picture(src, width=Cm(14))

                    # Caption
                    cap_text = comp.caption or comp.chart_title
                    if cap_text:
                        cp = doc.add_paragraph(cap_text)
                        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cr = cp.runs[0] if cp.runs else cp.add_run(cap_text)
                        cr.font.size = Pt(9)
                        cr.font.italic = True
                        cr.font.color.rgb = MGRAY
                except Exception as e:
                    logger.warning(f'Could not embed image in DOCX: {e}')
                    p = doc.add_paragraph(f'[Image: {comp.caption or comp.chart_title or "chart"}]')

        doc.add_paragraph()

    # ── Approval Chain ────────────────────────────────────────────
    approval_steps = list(sop.approval_chain.all())
    if approval_steps:
        ah = doc.add_heading('Approval Chain', level=1)
        ah.runs[0].font.color.rgb = BLUE

        at = doc.add_table(rows=1 + len(approval_steps), cols=5)
        at.style = 'Table Grid'
        headers = ['Step', 'Role', 'Approver', 'Status', 'Date']
        for ci, h in enumerate(headers):
            cell = at.cell(0, ci)
            set_cell_bg(cell, '003DA5')
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)

        for ri, step in enumerate(approval_steps):
            row_data = [
                str(step.step),
                step.role,
                step.approver.get_full_name() if step.approver else '—',
                step.get_status_display(),
                step.timestamp.strftime('%d %b %Y') if step.timestamp else 'Pending',
            ]
            for ci, text in enumerate(row_data):
                cell = at.cell(ri + 1, ci)
                p = cell.paragraphs[0]
                run = p.add_run(str(text))
                run.font.size = Pt(9)
                run.font.color.rgb = DGRAY
                if ri % 2 == 0:
                    set_cell_bg(cell, 'F2F4F8')
        doc.add_paragraph()

    # ── Footer ────────────────────────────────────────────────────
    fp = doc.add_paragraph()
    add_horizontal_line(fp, 'D0D7E3', 4)
    footer_run = fp.add_run(
        f'Generated: {datetime.now().strftime("%d %b %Y %H:%M")} | HPCL SOP Management Portal'
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = MGRAY
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
